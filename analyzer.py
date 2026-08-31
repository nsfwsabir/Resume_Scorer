from pathlib import Path
from groq import Groq
from pydantic import BaseModel,Field
from dotenv import load_dotenv
import json
import os
import time
import sys
from pypdf import PdfReader
from docx import Document


load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API Key not found!")

client=Groq(api_key=my_api_key)
model="openai/gpt-oss-120b"

print("Paste the JOB Description.")
print("Type END on a new line when finished.\n")
job_description= ""
while True:
    line=input()
    if line.strip().upper() == "END":
        break
    job_description += line + "\n"


class JobDesc(BaseModel):
    role:str
    required_skills:list[str]
    preferred_skils:list[str]
    minimum_experience: float | None
    educational_requirements: list[str]
    responsibilities: list[str]


jobdesc_schema= JobDesc.model_json_schema()


system_prompt= f"""
    You are an expert HR Assisstant

    Your job is to analyze job description and extract structured information from them.

    Return ONLY valid JSON matching this schema: {jobdesc_schema}

    IMPORTANT:
    DO NOT return the scheme itself.
    DO NOT return fields like "properties", "title" or "type".
    Fill the schema with actual information extracted from the job description.

    If minimum experience is not mentioned, return null.
    If information for a list is missing, return an empty list.
    DO NOT invent information. 

    """

user_prompt= f"""
    Analyze the following job description.

    {job_description}
    """

message_system= {
    "role": "system",
    "content": system_prompt
    }
message_user={
    "role": "user",
    "content": user_prompt
    }
response_format={
    "type":"json_object"
    }

messages=[message_system, message_user]
response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
answer=response.choices[0].message.content


job_data=json.loads(answer)
job=JobDesc(**job_data)

class MatchResult(BaseModel):
    score: float
    candidate_name: str | None = None
    matching_skills: list[str] = []
    missing_skills: list[str] = []
    experience_requirement_met: bool
    verdict: str
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None
    skills:list[str]= Field(default_factory=list)
    experiences:list[Experience]= Field(default_factory=list)
    education:list[str]= Field(default_factory=list)
    projects:list[str]= Field(default_factory=list)
    certifications:list[str]= Field(default_factory=list)

resume_schema= Resume.model_json_schema()


def read_pdf(file_path):
    reader= PdfReader(file_path)
    text=""
    for page in reader.pages:
        page_text= page.extract_text()
        if page_text:
            text+=page_text + "\n"
    return text

def read_docx(file_path):
    document= Document(file_path)
    text=""
    for para in document.paragraphs:
        if para.text.strip():
            text+=para.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text+=cell.text + "\n"
    return text

def read_resume(file_path):
    if file_path.suffix.lower() ==".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() ==".docx":
        return read_docx(file_path)
    else:
        return None 

def parse_resume(resume_text):
    system_prompt=f"""
        You are an expert parser.

        Extract information from the resume based on its meaning,
        not only based on exact section headings.

        Different resumes may use different headings.

        For example:
        - Experience
        - Professional Experience
        - Work History
        - Employment 
        - Internships

        They may all content relevant experience.

        Skills may also appear in the skills section, work experience, 
        internships or projects.

        Return ONLY valid JSON mathching this schema: 

        {resume_schema}

        Important rules:

        1. DO NOT invent information.
        2. If a value is not available, return null.
        3. If a list has not information, return an empty list.
        4. Internships inside experience.
        5. Extract skills mentioned across the entire resume.    
    
        """

    user_prompt=f"""
        Parse the following resume:

        {resume_text}    
    
        """
    
    message_system= {
    "role": "system",
    "content": system_prompt
    }
    message_user={
    "role": "user",
    "content": user_prompt
    }
    response_format={
    "type":"json_object"
    }
    messages=[message_system, message_user]
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output=response.choices[0].message.content
    data=json.loads(raw_output)
    resume= Resume(**data)
    return resume



def final_score(job,resume):
    match_schema= MatchResult.model_json_schema()
    prompt= f"""
        You are an HR Recruiter.

        Compare the candidate's resume with the job description.

        JOB DESCRIPTION:
        {job.model_dump_json(indent=2)}

        CANDIDATE'S RESUME:
        {resume.model_dump_json(indent=2)}

        Return JSON matching this schema:

        {match_schema}

        Give me:

        1. Candidate's Name.
        2. Matching skills.
        3. Missing important skills.
        4. Wether experience requirement is met.
        5. Overall match percentage ranging 0 to 100.
        6. A short final verdict.

        Keep the response concise and easy to read.
        """
    message={
    "role": "user",
    "content": prompt
    }
    response_format={
    "type":"json_object"
    }
    messages=[message]
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data=json.loads(response.choices[0].message.content)
    return MatchResult(**data)


resume_folder=Path("resumes")
all_results=[]

for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\n Processing:", file_path.name)
    resume_text=read_resume(file_path)
    parsed_resume=parse_resume(resume_text)
    time.sleep(5)
    result=final_score(job,parsed_resume)
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "matching_skills": result.matching_skills,
        "missing_skills": result.missing_skills,
        "experience_requirement_met": result.experience_requirement_met,
        "verdict": result.verdict
    })
all_results.sort(
        key=lambda candidate: candidate["score"],
        reverse=True
    )
top_2=all_results[:2]
if len(all_results) <=2:
    bottom_2=[]
else:
    bottom_2=all_results[-2:]

print("\n===== TOP 2 CANDIDATES =====")

for candidate in top_2:
    print(f"\n{candidate['name']} - {candidate['score']}%")
    print("Matching skills:", candidate["matching_skills"])
    print("Missing skills:", candidate["missing_skills"])
    print(
        "Experience requirement met:",
        candidate["experience_requirement_met"]
    )
    print("Verdict:", candidate["verdict"])


print("\n===== BOTTOM 2 CANDIDATES =====")

for candidate in bottom_2:
    print(f"\n{candidate['name']} - {candidate['score']}%")
    print("Matching skills:", candidate["matching_skills"])
    print("Missing skills:", candidate["missing_skills"])
    print(
        "Experience requirement met:",
        candidate["experience_requirement_met"]
    )
    print("Verdict:", candidate["verdict"])

